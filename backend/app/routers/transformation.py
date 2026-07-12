import html
import os, shutil, uuid
import re
import tempfile
from pathlib import Path
from fastapi import APIRouter, UploadFile, File, Form, HTTPException
from fastapi.responses import FileResponse
from openpyxl import Workbook, load_workbook
from openpyxl.cell.cell import MergedCell
from openpyxl.utils import get_column_letter
from docx import Document
import pdfplumber
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfgen import canvas
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

from app.services.transformation.extractors.excel_extractor import extract_excel
from app.services.transformation.template_engine.template_scanner import scan_template
from app.services.transformation.template_engine.injector import inject_into_template
from app.services.transformation.mappers.semantic_mapper import semantic_map
from app.services.transformation.exporters.word_writer import excel_to_word
from app.services.transformation.exporters.pdf_writer import excel_to_pdf

router = APIRouter(prefix="/transform", tags=["Transformation"])

UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)
CONVERTER_DIR = Path(tempfile.gettempdir()) / "teachassist_converted"
CONVERTER_DIR.mkdir(parents=True, exist_ok=True)

@router.post("/")
async def transform(
    source_file: UploadFile = File(...),
    template_file: UploadFile = File(...),
    output_type: str = Form(...)
):
    uid = uuid.uuid4().hex

    src_path = os.path.join(UPLOAD_DIR, f"src_{uid}_{source_file.filename}")
    tmp_path = os.path.join(UPLOAD_DIR, f"tmp_{uid}_{template_file.filename}")

    with open(src_path, "wb") as f:
        shutil.copyfileobj(source_file.file, f)
    with open(tmp_path, "wb") as f:
        shutil.copyfileobj(template_file.file, f)

    source_df = extract_excel(src_path)
    template = scan_template(tmp_path)

    mapping = semantic_map(source_df, template["fields"])
    if not mapping:
        raise Exception("No mapping found")

    temp_excel = os.path.join(UPLOAD_DIR, f"temp_{uid}.xlsx")

    inject_into_template(
        source_df, mapping, tmp_path, temp_excel, "xlsx"
    )

    if output_type == "word":
        final = os.path.join(UPLOAD_DIR, f"output_{uid}.docx")
        excel_to_word(temp_excel, final)
    elif output_type == "pdf":
        final = os.path.join(UPLOAD_DIR, f"output_{uid}.pdf")
        excel_to_pdf(temp_excel, final)
    else:
        final = temp_excel

    return {"file": os.path.basename(final)}

@router.get("/uploads/{filename}")
def download(filename: str):
    return FileResponse(os.path.join(UPLOAD_DIR, filename))


@router.post("/convert-document")
async def convert_document(
    file: UploadFile = File(...),
    conversion_type: str = Form(...),
):
    suffix = Path(file.filename or "").suffix.lower()
    uid = uuid.uuid4().hex
    input_path = CONVERTER_DIR / f"input_{uid}{suffix}"
    with input_path.open("wb") as handle:
        shutil.copyfileobj(file.file, handle)

    try:
        if conversion_type == "excel_to_pdf":
            if suffix not in {".xlsx", ".xlsm", ".xls"}:
                raise HTTPException(status_code=400, detail="Please upload an Excel file.")
            output_path = CONVERTER_DIR / f"converted_{uid}.pdf"
            convert_excel_to_pdf(input_path, output_path)
            return FileResponse(output_path, filename=f"{Path(file.filename or 'excel').stem}.pdf", media_type="application/pdf")

        if conversion_type == "word_to_pdf":
            if suffix not in {".docx"}:
                raise HTTPException(status_code=400, detail="Please upload a DOCX Word file.")
            output_path = CONVERTER_DIR / f"converted_{uid}.pdf"
            convert_word_to_pdf(input_path, output_path)
            return FileResponse(output_path, filename=f"{Path(file.filename or 'word').stem}.pdf", media_type="application/pdf")

        if conversion_type == "pdf_to_excel":
            if suffix != ".pdf":
                raise HTTPException(status_code=400, detail="Please upload a PDF file.")
            output_path = CONVERTER_DIR / f"converted_{uid}.xlsx"
            convert_pdf_to_excel(input_path, output_path)
            return FileResponse(
                output_path,
                filename=f"{Path(file.filename or 'pdf').stem}.xlsx",
                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
            )

        raise HTTPException(status_code=400, detail="Please choose a valid conversion.")
    except HTTPException:
        raise
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Conversion failed: {exc}") from exc


def convert_excel_to_pdf(input_path: Path, output_path: Path) -> None:
    workbook = load_workbook(input_path, data_only=True)
    page_size = landscape(A4)
    pdf = canvas.Canvas(str(output_path), pagesize=page_size)
    page_width, page_height = page_size
    left_margin = 18
    right_margin = 18
    top_margin = 18
    bottom_margin = 18
    available_width = page_width - left_margin - right_margin

    # Convert only the active worksheet. Mapping outputs can contain helper
    # sheets like "Unmapped Columns"; teachers expect the visible main sheet.
    for sheet in [workbook.active]:
        bounds = used_sheet_bounds(sheet)
        if not bounds:
            continue

        min_row, max_row, min_col, max_col = bounds
        max_visible_cols = min(max_col, min_col + 17)
        col_widths = excel_column_widths(sheet, min_col, max_visible_cols, available_width)
        y = page_height - top_margin
        for row in range(min_row, max_row + 1):
            row_height = excel_pdf_row_height(sheet, row, min_col, max_visible_cols, col_widths)
            if y - row_height < bottom_margin:
                pdf.showPage()
                y = page_height - top_margin
            draw_excel_pdf_row(pdf, sheet, row, min_col, max_visible_cols, col_widths, left_margin, y, row_height)
            y -= row_height

    pdf.save()


def excel_cell_text(value) -> str:
    if value is None:
        return ""
    text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f]", "", str(value))
    return text


def excel_pdf_row_height(sheet, row: int, min_col: int, max_col: int, col_widths: list[float]) -> float:
    configured = sheet.row_dimensions[row].height
    base = float(configured) if configured else 20.0
    max_lines = 1
    for offset, col in enumerate(range(min_col, max_col + 1)):
        cell = resolved_excel_cell(sheet, row, col)
        text = "" if is_merged_continuation(sheet, row, col) else excel_cell_text(cell.value)
        max_lines = max(max_lines, len(wrap_pdf_text(text, col_widths[offset] - 6, 6.2)))
    return max(base, min(88.0, max_lines * 7.5 + 8))


def draw_excel_pdf_row(pdf, sheet, row: int, min_col: int, max_col: int, col_widths: list[float], x_start: float, y_top: float, row_height: float) -> None:
    x = x_start
    for offset, col in enumerate(range(min_col, max_col + 1)):
        width = col_widths[offset]
        cell = resolved_excel_cell(sheet, row, col)
        bg = excel_cell_background(cell) or colors.white
        pdf.setFillColor(bg)
        pdf.rect(x, y_top - row_height, width, row_height, fill=1, stroke=0)
        pdf.setStrokeColor(colors.lightgrey)
        pdf.setLineWidth(0.25)
        pdf.rect(x, y_top - row_height, width, row_height, fill=0, stroke=1)

        text = "" if is_merged_continuation(sheet, row, col) else excel_cell_text(cell.value)
        pdf.setFillColor(colors.black)
        pdf.setFont("Helvetica-Bold" if cell.font and cell.font.bold else "Helvetica", 6.2)
        lines = wrap_pdf_text(text, width - 6, 6.2)
        text_y = y_top - 8
        horizontal = (cell.alignment.horizontal or "").lower() if cell.alignment else ""
        for line in lines[: max(1, int((row_height - 4) / 7.5))]:
            draw_x = x + 3
            if horizontal in {"center", "centercontinuous"}:
                draw_x = x + max(3, (width - pdf.stringWidth(line, pdf._fontname, 6.2)) / 2)
            elif horizontal == "right":
                draw_x = x + max(3, width - pdf.stringWidth(line, pdf._fontname, 6.2) - 3)
            pdf.drawString(draw_x, text_y, line)
            text_y -= 7.5
        x += width


def wrap_pdf_text(text: str, width: float, font_size: float) -> list[str]:
    if not text:
        return [""]
    chars_per_line = max(4, int(width / (font_size * 0.52)))
    lines: list[str] = []
    for raw_line in str(text).splitlines() or [""]:
        words = raw_line.split()
        if not words:
            lines.append("")
            continue
        current = ""
        for word in words:
            candidate = f"{current} {word}".strip()
            if len(candidate) <= chars_per_line:
                current = candidate
            else:
                if current:
                    lines.append(current)
                while len(word) > chars_per_line:
                    lines.append(word[:chars_per_line])
                    word = word[chars_per_line:]
                current = word
        if current:
            lines.append(current)
    return lines or [""]


def used_sheet_bounds(sheet) -> tuple[int, int, int, int] | None:
    min_row = min_col = None
    max_row = max_col = 0
    for row in sheet.iter_rows():
        for cell in row:
            if cell.value in (None, ""):
                continue
            min_row = cell.row if min_row is None else min(min_row, cell.row)
            min_col = cell.column if min_col is None else min(min_col, cell.column)
            max_row = max(max_row, cell.row)
            max_col = max(max_col, cell.column)
    if min_row is None or min_col is None:
        return None
    return min_row, min(max_row, min_row + 150), min_col, min(max_col, min_col + 25)


def excel_table_styles(sheet, table_rows: list[int], min_col: int, max_col: int) -> list[tuple]:
    style_commands = [
        ("GRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ("FONTSIZE", (0, 0), (-1, -1), 6.5),
    ]
    for table_row_index, row in enumerate(table_rows):
        for col in range(min_col, max_col + 1):
            cell = sheet.cell(row=row, column=col)
            table_pos = (col - min_col, table_row_index)
            if cell.font and cell.font.bold:
                style_commands.append(("FONTNAME", table_pos, table_pos, "Helvetica-Bold"))
            bg = excel_cell_background(cell)
            if bg:
                style_commands.append(("BACKGROUND", table_pos, table_pos, bg))
            horizontal = (cell.alignment.horizontal or "").lower() if cell.alignment else ""
            if horizontal in {"center", "centercontinuous"}:
                style_commands.append(("ALIGN", table_pos, table_pos, "CENTER"))
            elif horizontal == "right":
                style_commands.append(("ALIGN", table_pos, table_pos, "RIGHT"))
    return style_commands


def excel_column_widths(sheet, min_col: int, max_col: int, available_width: float) -> list[float]:
    raw_widths = []
    for col in range(min_col, max_col + 1):
        letter = get_column_letter(col)
        excel_width = sheet.column_dimensions[letter].width or 10
        raw_widths.append(max(30, min(90, float(excel_width) * 5.2)))
    total = sum(raw_widths) or 1
    if total <= available_width:
        return raw_widths
    scale = available_width / total
    return [max(26, width * scale) for width in raw_widths]


def resolved_excel_cell(sheet, row: int, col: int):
    cell = sheet.cell(row=row, column=col)
    if not isinstance(cell, MergedCell):
        return cell
    for merged_range in sheet.merged_cells.ranges:
        if cell.coordinate in merged_range:
            return sheet.cell(row=merged_range.min_row, column=merged_range.min_col)
    return cell


def is_merged_continuation(sheet, row: int, col: int) -> bool:
    coordinate = sheet.cell(row=row, column=col).coordinate
    for merged_range in sheet.merged_cells.ranges:
        if coordinate in merged_range:
            return row != merged_range.min_row or col != merged_range.min_col
    return False


def excel_cell_background(cell):
    try:
        fg = cell.fill.fgColor
        if not fg or fg.type != "rgb" or not fg.rgb:
            return None
        value = str(fg.rgb)[-6:]
        if value.upper() in {"000000", "FFFFFF"}:
            return None
        return colors.HexColor(f"#{value}")
    except Exception:
        return None


def convert_word_to_pdf(input_path: Path, output_path: Path) -> None:
    document = Document(str(input_path))
    pdf = SimpleDocTemplate(str(output_path), pagesize=A4, leftMargin=42, rightMargin=42, topMargin=42, bottomMargin=42)
    styles = getSampleStyleSheet()
    story = []

    for block in iter_word_blocks(document):
        if block.__class__.__name__ == "Paragraph":
            text = paragraph_to_html(block)
            if not text.strip():
                story.append(Spacer(1, 6))
                continue
            style = paragraph_style(block, styles)
            story.append(Paragraph(text, style))
            story.append(Spacer(1, 4))
        else:
            table_data = []
            for row in block.rows:
                table_data.append([paragraph_to_html(cell.paragraphs[0]) if cell.paragraphs else "" for cell in row.cells])
            if table_data:
                table = Table(table_data, repeatRows=1)
                table.setStyle(TableStyle([
                    ("GRID", (0, 0), (-1, -1), 0.3, colors.lightgrey),
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#eef3ff")),
                    ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                    ("VALIGN", (0, 0), (-1, -1), "TOP"),
                    ("FONTSIZE", (0, 0), (-1, -1), 8),
                ]))
                story.append(table)
                story.append(Spacer(1, 8))

    pdf.build(story)


def iter_word_blocks(document: Document):
    body = document.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            for paragraph in document.paragraphs:
                if paragraph._p is child:
                    yield paragraph
                    break
        elif child.tag.endswith("}tbl"):
            for table in document.tables:
                if table._tbl is child:
                    yield table
                    break


def paragraph_to_html(paragraph) -> str:
    pieces = []
    for run in paragraph.runs:
        text = html.escape(run.text or "")
        if not text:
            continue
        if run.bold:
            text = f"<b>{text}</b>"
        if run.italic:
            text = f"<i>{text}</i>"
        if run.underline:
            text = f"<u>{text}</u>"
        pieces.append(text)
    return "".join(pieces) or html.escape(paragraph.text or "")


def paragraph_style(paragraph, styles):
    name = (paragraph.style.name or "").lower() if paragraph.style else ""
    if "heading 1" in name:
        return styles["Heading1"]
    if "heading 2" in name:
        return styles["Heading2"]
    alignment = paragraph.alignment
    base = ParagraphStyle("DocParagraph", parent=styles["Normal"], leading=14, spaceAfter=3)
    if alignment == 1:
        base.alignment = TA_CENTER
    elif alignment == 2:
        base.alignment = TA_RIGHT
    return base


def convert_pdf_to_excel(input_path: Path, output_path: Path) -> None:
    workbook = Workbook()
    workbook.remove(workbook.active)
    with pdfplumber.open(str(input_path)) as pdf:
        for page_index, page in enumerate(pdf.pages, start=1):
            sheet = workbook.create_sheet(title=f"Page {page_index}")
            tables = page.extract_tables() or []
            row_cursor = 1
            if tables:
                for table in tables:
                    for row in table:
                        for col_index, value in enumerate(row or [], start=1):
                            sheet.cell(row=row_cursor, column=col_index).value = value
                        row_cursor += 1
                    row_cursor += 1
            else:
                text = page.extract_text() or ""
                for line in text.splitlines():
                    sheet.cell(row=row_cursor, column=1).value = line
                    row_cursor += 1
    if not workbook.sheetnames:
        workbook.create_sheet(title="PDF")
    workbook.save(output_path)
