"""
app/services/cis_parser.py

CIS (Course Information Sheet) se weekly topics AND CLO taxonomy levels extract karta hai.
PDF aur Word (.docx) dono support karta hai.
"""

import io
import re
from typing import List, Dict, Any
from fastapi import UploadFile


async def parse_cis(file: UploadFile) -> List[Dict[str, Any]]:
    """
    CIS file parse karke weekly topics return karta hai.

    Returns:
        [ { "week": 1, "topics": "Intensity Transformation, Filters..." }, ... ]
    """
    filename = file.filename.lower()
    content  = await file.read()

    if filename.endswith(".pdf"):
        return _parse_pdf(content)
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        return _parse_docx(content)
    else:
        raise ValueError(f"Unsupported file type: {filename}. Use PDF or DOCX.")


def _parse_pdf(content: bytes) -> List[Dict[str, Any]]:
    import fitz
    doc  = fitz.open(stream=content, filetype="pdf")
    text = ""
    for page in doc:
        text += page.get_text()
    doc.close()
    print("=== CIS EXTRACTED TEXT ===")
    print(text[:1000])
    print("==========================")
    return _extract_weeks_from_text(text)


def _parse_docx(content: bytes) -> List[Dict[str, Any]]:
    from docx import Document
    doc   = Document(io.BytesIO(content))
    lines = []
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(
                cell.text.strip() for cell in row.cells if cell.text.strip()
            )
            if row_text:
                lines.append(row_text)
    for para in doc.paragraphs:
        if para.text.strip():
            lines.append(para.text.strip())
    text = "\n".join(lines)
    print("=== CIS EXTRACTED TEXT ===")
    print(text[:1000])
    print("==========================")
    return _extract_weeks_from_text(text)


def _extract_weeks_from_text(text: str) -> List[Dict[str, Any]]:
    weeks = []
    patterns = [
        r'(?:Week\s*)?(\d+)\s*\|?\s*([A-Z][^\n]{20,})',
        r'^(\d{1,2})\s+(.{20,})$',
    ]
    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        for pattern in patterns:
            match = re.match(pattern, line, re.IGNORECASE)
            if match:
                week_no = int(match.group(1))
                topic   = match.group(2).strip()
                if 1 <= week_no <= 18 and len(topic) > 15:
                    existing = next((w for w in weeks if w["week"] == week_no), None)
                    if existing:
                        existing["topics"] += " " + topic
                    else:
                        weeks.append({"week": week_no, "topics": topic})
                break
    weeks.sort(key=lambda x: x["week"])
    print(f"CIS se {len(weeks)} weeks extract hue")
    for w in weeks:
        print(f"  Week {w['week']}: {w['topics'][:80]}...")
    return weeks


# ═══════════════════════════════════════════════════════════
# NEW: CLO Taxonomy Extraction
# ═══════════════════════════════════════════════════════════

def extract_clo_taxonomy(text: str) -> Dict[str, str]:
    """
    CIS text se CLO → taxonomy level (C1-C6) mapping nikalo.

    CIS mein aksar yeh format hota hai:
      CLO 1 | Explain fundamentals... | C1 | ...
      CLO-2 | Describe techniques...  | C2 | ...
      1. | Analyze methods...         | C3 | ...

    Returns:
        { "CLO-1": "C1", "CLO-2": "C2", "CLO-3": "C3" }
    """
    clo_taxonomy: Dict[str, str] = {}

    lines = text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Pattern 1: "CLO-1 ... C1" or "CLO 1 ... C1" anywhere in line
        # Handles table rows like: "1. | Explain ... | C1 | ..."
        clo_match = re.search(
            r'(?:CLO[\s\-\.]?)?(\d+)[^\n]*?\b(C[1-6])\b',
            line, re.IGNORECASE
        )
        if clo_match:
            clo_num  = clo_match.group(1)
            taxonomy = clo_match.group(2).upper()
            key      = f"CLO-{clo_num}"
            if key not in clo_taxonomy:
                clo_taxonomy[key] = taxonomy
                print(f"Taxonomy found: {key} -> {taxonomy} | Line: {line[:80]}")
            continue

        # Pattern 2: pipe-separated table row
        # e.g. "CLO No. | CLO Description | Domain | Mapped PLO | Level"
        # e.g. "1 | Explain fundamentals | C1 | PLO1 | 3"
        if '|' in line:
            parts = [p.strip() for p in line.split('|')]
            # Find C1-C6 in any cell
            c_match = None
            for part in parts:
                m = re.fullmatch(r'C[1-6]', part, re.IGNORECASE)
                if m:
                    c_match = m.group(0).upper()
                    break
            if c_match:
                # Find CLO number — first cell that's a digit or starts with CLO
                for part in parts:
                    num_m = re.match(r'(?:CLO[\s\-\.]?)?(\d+)$', part.strip(), re.IGNORECASE)
                    if num_m:
                        key = f"CLO-{num_m.group(1)}"
                        if key not in clo_taxonomy:
                            clo_taxonomy[key] = c_match
                            print(f"Taxonomy (table): {key} -> {c_match} | Line: {line[:80]}")
                        break

    if not clo_taxonomy:
        print("No CLO taxonomy found in CIS - will fallback to CLO number mapping")

    return clo_taxonomy


def extract_course_title(text: str) -> str:
    patterns = [
        r'Course\s*(?:Title|Name)\s*[:\-]\s*(.+)',
        r'Subject\s*(?:Title|Name)?\s*[:\-]\s*(.+)'
    ]
    for line in text.split('\n'):
        clean = line.strip()
        if not clean:
            continue
        for pattern in patterns:
            match = re.search(pattern, clean, re.IGNORECASE)
            if match:
                return match.group(1).strip()
    return ""


async def parse_cis_full(file: UploadFile) -> Dict[str, Any]:
    """
    CIS file se BOTH weekly topics AND CLO taxonomy extract karta hai.

    Returns:
        {
          "weeks":        [ { "week": 1, "topics": "..." }, ... ],
          "clo_taxonomy": { "CLO-1": "C1", "CLO-2": "C2", "CLO-3": "C3" }
        }
    """
    filename = file.filename.lower()
    content  = await file.read()

    if filename.endswith(".pdf"):
        import fitz
        doc  = fitz.open(stream=content, filetype="pdf")
        text = ""
        for page in doc:
            text += page.get_text()
        doc.close()
    elif filename.endswith(".docx") or filename.endswith(".doc"):
        from docx import Document
        doc   = Document(io.BytesIO(content))
        lines = []
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        text = "\n".join(lines)
    else:
        raise ValueError(f"Unsupported file type: {filename}")

    weeks        = _extract_weeks_from_text(text)
    clo_taxonomy = extract_clo_taxonomy(text)
    course_title = extract_course_title(text)

    return {
        "weeks":        weeks,
        "clo_taxonomy": clo_taxonomy,
        "course_title": course_title
    }


def format_cis_for_prompt(cis_weeks: List[Dict[str, Any]]) -> str:
    if not cis_weeks:
        return "  (No CIS data available)"
    lines = []
    for w in cis_weeks:
        lines.append(f"  Week {w['week']}: {w['topics']}")
    return "\n".join(lines)
